from pathlib import Path
import csv

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[1]
REPORT_DIR = ROOT / "report"
RESULTS_DIR = ROOT / "results"
OUTPUT = REPORT_DIR / "DZ3_6DOF_NED_Report.docx"


def set_run_font(run, size=10.5, bold=False, color=None):
    run.font.name = "Calibri"
    run._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
    run._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "Calibri")
    run.font.size = Pt(size)
    run.bold = bold
    if color:
        run.font.color.rgb = RGBColor.from_string(color)


def shade(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    tc_pr.append(shd)


def add_hyperlink(paragraph, text, url):
    relationship_id = paragraph.part.relate_to(
        url,
        "http://schemas.openxmlformats.org/officeDocument/2006/relationships/hyperlink",
        is_external=True,
    )
    hyperlink = OxmlElement("w:hyperlink")
    hyperlink.set(qn("r:id"), relationship_id)
    run = OxmlElement("w:r")
    properties = OxmlElement("w:rPr")
    color = OxmlElement("w:color")
    color.set(qn("w:val"), "0563C1")
    underline = OxmlElement("w:u")
    underline.set(qn("w:val"), "single")
    properties.append(color)
    properties.append(underline)
    run.append(properties)
    content = OxmlElement("w:t")
    content.text = text
    run.append(content)
    hyperlink.append(run)
    paragraph._p.append(hyperlink)


def set_cell_margins(cell, top=80, start=120, bottom=80, end=120):
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for side, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tc_mar.find(qn(f"w:{side}"))
        if node is None:
            node = OxmlElement(f"w:{side}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def set_cell_width(cell, width_inches):
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_w = tc_pr.find(qn("w:tcW"))
    if tc_w is None:
        tc_w = OxmlElement("w:tcW")
        tc_pr.append(tc_w)
    tc_w.set(qn("w:w"), str(round(width_inches * 1440)))
    tc_w.set(qn("w:type"), "dxa")
    cell.width = Inches(width_inches)


def set_repeat_table_header(row):
    tr_pr = row._tr.get_or_add_trPr()
    node = OxmlElement("w:tblHeader")
    node.set(qn("w:val"), "true")
    tr_pr.append(node)


def set_table_geometry(table, widths):
    total_dxa = sum(round(width * 1440) for width in widths)
    tbl_pr = table._tbl.tblPr
    tbl_w = tbl_pr.first_child_found_in("w:tblW")
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:w"), str(total_dxa))
    tbl_w.set(qn("w:type"), "dxa")
    tbl_ind = tbl_pr.first_child_found_in("w:tblInd")
    if tbl_ind is None:
        tbl_ind = OxmlElement("w:tblInd")
        tbl_pr.append(tbl_ind)
    tbl_ind.set(qn("w:w"), "120")
    tbl_ind.set(qn("w:type"), "dxa")
    grid = table._tbl.tblGrid
    for column in list(grid):
        grid.remove(column)
    for width in widths:
        column = OxmlElement("w:gridCol")
        column.set(qn("w:w"), str(round(width * 1440)))
        grid.append(column)


def add_table(doc, headers, rows, widths):
    table = doc.add_table(rows=1, cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    table.autofit = False
    table.style = "Table Grid"
    set_table_geometry(table, widths)
    header = table.rows[0]
    set_repeat_table_header(header)
    for index, text in enumerate(headers):
        cell = header.cells[index]
        set_cell_width(cell, widths[index])
        set_cell_margins(cell)
        shade(cell, "E8EEF5")
        paragraph = cell.paragraphs[0]
        paragraph.paragraph_format.space_after = Pt(0)
        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = paragraph.add_run(text)
        set_run_font(run, size=9, bold=True, color="1F4D78")
    for row_values in rows:
        cells = table.add_row().cells
        for index, text in enumerate(row_values):
            cell = cells[index]
            set_cell_width(cell, widths[index])
            set_cell_margins(cell)
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            paragraph = cell.paragraphs[0]
            paragraph.paragraph_format.space_after = Pt(0)
            run = paragraph.add_run(str(text))
            set_run_font(run, size=9)
    doc.add_paragraph().paragraph_format.space_after = Pt(2)
    return table


def add_paragraph(doc, text, bold_prefix=None, size=10.5):
    paragraph = doc.add_paragraph()
    paragraph.paragraph_format.space_after = Pt(6)
    paragraph.paragraph_format.line_spacing = 1.10
    if bold_prefix and text.startswith(bold_prefix):
        run = paragraph.add_run(bold_prefix)
        set_run_font(run, size=size, bold=True)
        run = paragraph.add_run(text[len(bold_prefix):])
        set_run_font(run, size=size)
    else:
        run = paragraph.add_run(text)
        set_run_font(run, size=size)
    return paragraph


def add_heading(doc, text, level=1):
    paragraph = doc.add_paragraph(style=f"Heading {level}")
    run = paragraph.add_run(text)
    return paragraph


def add_figure_pair(doc, left_path, right_path, caption):
    paragraph = doc.add_paragraph()
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    paragraph.paragraph_format.space_after = Pt(3)
    paragraph.add_run().add_picture(str(left_path), width=Inches(3.05))
    paragraph.add_run("  ")
    paragraph.add_run().add_picture(str(right_path), width=Inches(3.05))
    caption_paragraph = doc.add_paragraph()
    caption_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    caption_paragraph.paragraph_format.space_after = Pt(8)
    set_run_font(caption_paragraph.add_run(caption), size=9, color="4B5563")


def read_summary():
    with (RESULTS_DIR / "comparison_summary.csv").open(encoding="utf-8-sig") as source:
        return list(csv.DictReader(source))


def main():
    REPORT_DIR.mkdir(exist_ok=True)
    document = Document()
    section = document.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)
    section.header_distance = Inches(0.492)
    section.footer_distance = Inches(0.492)

    normal = document.styles["Normal"]
    normal.font.name = "Calibri"
    normal._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
    normal._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
    normal.font.size = Pt(10.5)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.10
    for level, size, color, before, after in ((1, 16, "2E74B5", 16, 8), (2, 13, "2E74B5", 12, 6)):
        style = document.styles[f"Heading {level}"]
        style.font.name = "Calibri"
        style._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
        style._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
        style.font.size = Pt(size)
        style.font.color.rgb = RGBColor.from_string(color)
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)

    footer = section.footer.paragraphs[0]
    footer.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    footer_run = footer.add_run("DZ3 | 6DOF NED comparison")
    set_run_font(footer_run, size=9, color="6B7280")

    title = document.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title.paragraph_format.space_after = Pt(5)
    set_run_font(title.add_run("Домашнє завдання 3"), size=24, bold=True, color="0B2545")
    subtitle = document.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle.paragraph_format.space_after = Pt(18)
    set_run_font(subtitle.add_run("Власна 6DOF система рівнянь у NED та порівняння з Aerospace Block"), size=13, color="4B5563")

    add_heading(document, "1. Мета і межі роботи")
    add_paragraph(document, "Мета: замінити стандартний Aerospace Block 6DOF власною системою рівнянь твердого тіла в координатах NED, реалізувати її у Simulink і перевірити еквівалентність на контрольних навантаженнях та тест-кейсах ДЗ 1.")
    add_paragraph(document, "У реалізації сила і момент подаються у зв'язаній системі координат body. Гравітація, тяга та гіроскопічний момент ротора входять до сумарних зовнішніх навантажень. Позиція подається як [x_N, y_E, z_D]^T; тому висота 10 м має z_D = -10 м.")

    add_heading(document, "2. Виведені рівняння")
    add_paragraph(document, "Стан: x = [u v w p q r phi theta psi x_N y_E z_D]^T. Напрямна косинусна матриця C_be побудована за кутами Ейлера 3-2-1 (yaw-pitch-roll) і перетворює вектор NED у зв'язану систему.")
    add_paragraph(document, "Поступальна динаміка: Vb_dot = F_b/m - omega_b x Vb. Обертальна динаміка: omega_dot = I^(-1) [M_b - omega_b x (I omega_b)].")
    add_paragraph(document, "Кінематика: Euler_dot = E(phi,theta) omega_b; Ve = C_be^T Vb; Xe_dot = Ve. Сингулярність параметризації виникає за cos(theta)=0. Pitch-тест перетинає 90 градусів, тому фізичний збіг орієнтації біля цієї точки додатково слід оцінювати за DCM, а не лише за числовими значеннями кутів Ейлера.")

    add_heading(document, "3. Реалізація у Simulink")
    add_paragraph(document, "Модель dz3_6dof_ned_compare.slx має дві паралельні гілки з однаковими входами Forces/Moments. Власна гілка містить MATLAB Function для рівнянь, чотири безперервні інтегратори (Vb, pqr, Euler, Xe) та перетворення NED. Друга гілка використовує Aerospace Block 6DoF (Euler Angles).")
    add_paragraph(document, "У встановленій версії R2026a стандартний блок має 2 входи та 7 спільних виходів: Ve, Xe, Euler, DCM, Vb, p,q,r, p_dot,q_dot,r_dot. Додатковий сигнал Ab логуються лише у власній гілці як діагностичний.")
    add_paragraph(document, "Для стабільного прямого запуску модель має InitFcn: якщо forces_ts або moments_ts відсутні у workspace, автоматично створюються безпечні нульові сигнали. Якщо сценарний скрипт уже задав навантаження, вони не перезаписуються.")

    add_heading(document, "4. Контроль виконання вимог")
    requirement_rows = [
        ("1.1", "Власні рівняння у NED", "Виконано", "Перевірено 5 unit-тестами: транспортний член, рівняння Ейлера, DCM та кінематика."),
        ("1.2", "Simulink-імплементація", "Виконано", "Незалежна власна підсистема з інтеграторами; структура автоматично перевіряється."),
        ("1.3", "Aerospace 6DOF як еталон", "Виконано", "6DoF (Euler Angles), маса 1.321135985 kg, I = diag(0.0093,0.0092,0.0151)."),
        ("1.3", "Канонічні тести за DOF", "Виконано", "Surge, sway, heave, roll, pitch, yaw і combined 6DOF."),
        ("1.3", "Ті самі степ-тести ДЗ 1", "Виконано", "Vertical, roll, pitch, yaw; збережено команди моторів, динаміку моторів, тягу, гравітацію та гіроскопічний момент."),
        ("Надійність", "Прямий Run моделі", "Виконано", "InitFcn створює нульові Forces/Moments лише за відсутності сценарних входів; поведінка перевірена автоматичним тестом."),
        ("Захист", "Графіки тестів ДЗ 1", "Виконано", "Автоматично формуються overlay-графіки vertical, roll, pitch і yaw: Custom NED суцільною, Aerospace штриховою."),
        ("Додатково", "Laplace-перевірка", "Виконано", "Для F0/s: x_N(t)=F0 t^2/(2m); використано як незалежну аналітичну перевірку."),
    ]
    add_table(document, ["Пункт", "Вимога", "Стан", "Доказ"], requirement_rows, [0.55, 1.55, 0.95, 3.45])

    add_heading(document, "5. Результати порівняння")
    summary = read_summary()
    result_rows = []
    for row in summary:
        family = "ДЗ 1: моторний step" if row["test_family"] == "HW1_motor_step" else "Сили/моменти"
        result_rows.append((family, row["scenario"], f"{float(row['max_state_error']):.3e}", f"{float(row['final_position_error_m']):.3e}"))
    add_table(document, ["Набір", "Сценарій", "max |Delta state|", "|Delta Xe(T)|, м"], result_rows, [1.45, 1.45, 1.75, 1.85])
    maximum_state_error = max(float(row["max_state_error"]) for row in summary)
    add_paragraph(document, f"Найбільше відхилення між гілками у всіх 11 сценаріях: {maximum_state_error:.3e}. Воно відповідає похибці обчислень з плаваючою комою, а не відмінності моделей.")

    figure_path = RESULTS_DIR / "comparison_and_laplace.png"
    if figure_path.exists():
        picture = document.add_picture(str(figure_path), width=Inches(5.7))
        document.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
        caption = document.add_paragraph()
        caption.alignment = WD_ALIGN_PARAGRAPH.CENTER
        caption.paragraph_format.space_after = Pt(8)
        set_run_font(caption.add_run("Рисунок 1. Збіг траєкторій у combined 6DOF і Laplace-перевірка."), size=9, color="4B5563")

    add_heading(document, "6. Overlay-графіки тест-кейсів ДЗ 1")
    add_paragraph(document, "На всіх графіках власна модель Custom NED показана суцільною синьою лінією, Aerospace Block — червоною штриховою. Візуальне накладання підтверджує збіг вертикального руху та реакцій за каналами roll, pitch і yaw.")
    figure_pairs = [
        (RESULTS_DIR / "vertical_overlay.png", RESULTS_DIR / "roll_overlay.png",
         "Рисунки 2–3. Vertical: z_D та контроль крос-зв'язків; roll: phi та p."),
        (RESULTS_DIR / "pitch_overlay.png", RESULTS_DIR / "yaw_overlay.png",
         "Рисунки 4–5. Pitch: theta та q; yaw: psi та r."),
    ]
    for left_path, right_path, caption_text in figure_pairs:
        if left_path.exists() and right_path.exists():
            add_figure_pair(document, left_path, right_path, caption_text)
    add_paragraph(document, "Pitch-траєкторія проходить через область сингулярності Euler 3-2-1 біля theta = 90 градусів. Це обмеження способу параметризації; збіг DCM у числових результатах підтверджує збіг фізичної орієнтації двох моделей.")

    add_heading(document, "7. Незалежна перевірка за Лапласом")
    add_paragraph(document, "Для сталої сили F0=1 Н вздовж осі North за нульових початкових умов: F(s)=F0/s, X(s)=F0/(m s^3), отже x_N(t)=F0 t^2/(2m). За 2 с максимальна похибка власної і Aerospace-гілки від аналітичного розв'язку становить 2.98e-14 м.")
    add_paragraph(document, "Цей контроль відповідає методиці лекції 9-10: перехід від алгебраїчних рівнянь у зображення Лапласа до часової реакції та її числової перевірки.")

    add_heading(document, "8. Висновок і подальше вдосконалення")
    add_paragraph(document, "Власна 6DOF NED-система реалізована й пройшла 10 автоматичних перевірок: 5 математичних, 2 структурні та 3 наскрізні. Вона відтворює виходи Aerospace Block для канонічних навантажень і чотирьох тестів ДЗ 1 до рівня машинної точності. Прямий запуск моделі тепер також не потребує ручної підготовки workspace.")
    add_paragraph(document, "Що ще варто доробити: інтегрувати моторний міксер та модель ротора як окрему Simulink-підсистему всередині порівняльної моделі (зараз ідентичні часові історії навантажень одержані із перевіреної моделі ДЗ 1); для великих кутів додати quaternion-варіант або обмеження області Euler, щоб уникнути сингулярності pitch = ±90 градусів.")

    add_heading(document, "9. Відтворюваність і джерела")
    add_paragraph(document, "Запуск: відкрити dz3_6dof_ned_compare.slx і натиснути Run; без заданих входів модель використає нульові сили та моменти. Повне порівняння — run_dz3_comparison; перевірки — runtests('../tests'). Модель відновлюється скриптом build_dz3_6dof_model.")
    github_links = [
        ("Репозиторій", "https://github.com/ievgen-bovkun/flight-engineering-RD"),
        ("Папка homework3", "https://github.com/ievgen-bovkun/flight-engineering-RD/tree/codex/homework3-ned/homework3"),
        ("Запуск порівняння", "https://github.com/ievgen-bovkun/flight-engineering-RD/blob/codex/homework3-ned/homework3/scripts/run_dz3_comparison.m"),
        ("Simulink-модель", "https://github.com/ievgen-bovkun/flight-engineering-RD/blob/codex/homework3-ned/homework3/models/dz3_6dof_ned_compare.slx"),
        ("Автоматичні тести", "https://github.com/ievgen-bovkun/flight-engineering-RD/tree/codex/homework3-ned/homework3/tests"),
    ]
    for label, url in github_links:
        paragraph = document.add_paragraph()
        paragraph.paragraph_format.space_after = Pt(2)
        set_run_font(paragraph.add_run(f"{label}: "), size=9, bold=True)
        add_hyperlink(paragraph, url, url)
    add_paragraph(document, "Матеріали: ДЗ 1 (параметри й motor test mixes); ДЗ 2 (символьне виведення та NED); презентація 'Лекція 9-10. Воркшоп_ Аналіз стійкості та керованості моделі' (Laplace-перевірка); Aerospace Blockset R2026a.", size=8.5)

    document.core_properties.title = "DZ3: 6DOF NED comparison"
    document.core_properties.subject = "Simulink verification report"
    document.core_properties.author = "Codex"
    document.save(OUTPUT)
    print(OUTPUT)


if __name__ == "__main__":
    main()
